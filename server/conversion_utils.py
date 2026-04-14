#!/usr/bin/env python3
"""
File conversion utilities using Pandoc, PyMuPDF, and Pillow.
Supports conversions between: txt, md, docx, pdf, and images (png, jpg).
"""

import os
import subprocess
import sys
from pathlib import Path
from typing import Optional, List
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import io


class FileConverter:
    """Main converter class handling all format conversions."""
    
    SUPPORTED_FORMATS = {
        'txt': 'text',
        'md': 'markdown',
        'docx': 'docx',
        'pdf': 'pdf',
        'png': 'image',
        'jpg': 'image',
        'jpeg': 'image',
    }
    
    PANDOC_FORMAT_MAP = {
        'txt': 'plain',
        'md': 'markdown',
        'docx': 'docx',
    }
    
    @staticmethod
    def get_file_format(filepath: str) -> str:
        """Extract and validate file format from filepath."""
        ext = Path(filepath).suffix.lower().lstrip('.')
        if ext not in FileConverter.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}")
        return ext
    
    @staticmethod
    def txt_to_md(input_path: str, output_path: str) -> None:
        """Convert TXT to Markdown."""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    @staticmethod
    def txt_to_docx(input_path: str, output_path: str) -> None:
        """Convert TXT to DOCX using python-docx."""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = Document()
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        doc.save(output_path)
    
    @staticmethod
    def md_to_txt(input_path: str, output_path: str) -> None:
        """Convert Markdown to TXT (simple copy)."""
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    @staticmethod
    def md_to_docx(input_path: str, output_path: str) -> None:
        """Convert Markdown to DOCX using Pandoc."""
        try:
            subprocess.run(
                ['pandoc', input_path, '-o', output_path, '-f', 'markdown', '-t', 'docx'],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr.decode()}")
    
    @staticmethod
    def docx_to_txt(input_path: str, output_path: str) -> None:
        """Convert DOCX to TXT using python-docx."""
        doc = Document(input_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    @staticmethod
    def docx_to_md(input_path: str, output_path: str) -> None:
        """Convert DOCX to Markdown using Pandoc."""
        try:
            subprocess.run(
                ['pandoc', input_path, '-o', output_path, '-f', 'docx', '-t', 'markdown'],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr.decode()}")
    
    @staticmethod
    def txt_to_pdf(input_path: str, output_path: str) -> None:
        """Convert TXT to PDF using ReportLab (no external engine needed)."""
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch

        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read()

        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
                story.append(Spacer(1, 0.1 * inch))
            else:
                story.append(Spacer(1, 0.15 * inch))
        if not story:
            story.append(Paragraph(' ', styles['Normal']))
        doc.build(story)
    
    @staticmethod
    def md_to_pdf(input_path: str, output_path: str) -> None:
        """Convert Markdown to PDF using Pandoc with weasyprint."""
        try:
            subprocess.run(
                ['pandoc', input_path, '-o', output_path, '-f', 'markdown', '--pdf-engine=weasyprint'],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr.decode()}")
    
    @staticmethod
    def docx_to_pdf(input_path: str, output_path: str) -> None:
        """Convert DOCX to PDF using Pandoc with weasyprint."""
        try:
            subprocess.run(
                ['pandoc', input_path, '-o', output_path, '-f', 'docx', '--pdf-engine=weasyprint'],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"Pandoc conversion failed: {e.stderr.decode()}")
    
    @staticmethod
    def images_to_pdf(image_paths: List[str], output_path: str) -> None:
        """Convert one or more images to PDF using Pillow and ReportLab."""
        images = []
        for img_path in image_paths:
            img = Image.open(img_path)
            if img.mode in ('RGBA', 'LA', 'P'):
                # Convert to RGB for PDF compatibility
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                images.append(rgb_img.convert('RGB'))
            else:
                images.append(img.convert('RGB'))
        
        if images:
            images[0].save(output_path, save_all=True, append_images=images[1:])
    
    @staticmethod
    def pdf_to_images(input_path: str, output_dir: str) -> List[str]:
        """Convert PDF pages to images using PyMuPDF."""
        pdf_document = fitz.open(input_path)
        output_paths = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
            
            output_path = os.path.join(output_dir, f'page_{page_num + 1}.png')
            pix.save(output_path)
            output_paths.append(output_path)
        
        pdf_document.close()
        return output_paths
    
    @staticmethod
    def pdf_to_txt(input_path: str, output_path: str) -> None:
        """Convert PDF to TXT using PyMuPDF."""
        pdf_document = fitz.open(input_path)
        text = ""
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text += page.get_text()
        
        pdf_document.close()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    @staticmethod
    def pdf_to_md(input_path: str, output_path: str) -> None:
        """Convert PDF to Markdown using PyMuPDF (basic conversion)."""
        pdf_document = fitz.open(input_path)
        text = ""
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text += f"## Page {page_num + 1}\n\n"
            text += page.get_text()
            text += "\n\n---\n\n"
        
        pdf_document.close()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    @staticmethod
    def pdf_to_docx(input_path: str, output_path: str) -> None:
        """Convert PDF to DOCX using PyMuPDF and python-docx."""
        pdf_document = fitz.open(input_path)
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            
            if page_num > 0:
                doc.add_page_break()
            
            if text.strip():
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
        
        pdf_document.close()
        doc.save(output_path)
    
    @staticmethod
    def convert(input_path: str, output_path: str, input_format: Optional[str] = None, 
                output_format: Optional[str] = None) -> None:
        """
        Main conversion method that routes to appropriate converter.
        
        Args:
            input_path: Path to input file
            output_path: Path to output file
            input_format: Input format (auto-detected if not provided)
            output_format: Output format (auto-detected if not provided)
        """
        # Auto-detect formats if not provided
        if not input_format:
            input_format = FileConverter.get_file_format(input_path)
        if not output_format:
            output_format = FileConverter.get_file_format(output_path)
        
        input_format = input_format.lower().lstrip('.')
        output_format = output_format.lower().lstrip('.')
        
        # Validate formats
        if input_format not in FileConverter.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported input format: {input_format}")
        if output_format not in FileConverter.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported output format: {output_format}")
        
        # Same format - just copy
        if input_format == output_format:
            with open(input_path, 'rb') as src:
                with open(output_path, 'wb') as dst:
                    dst.write(src.read())
            return
        
        # Route to appropriate converter
        conversion_key = f"{input_format}_to_{output_format}"
        
        if conversion_key == "txt_to_md":
            FileConverter.txt_to_md(input_path, output_path)
        elif conversion_key == "txt_to_docx":
            FileConverter.txt_to_docx(input_path, output_path)
        elif conversion_key == "txt_to_pdf":
            FileConverter.txt_to_pdf(input_path, output_path)
        elif conversion_key == "md_to_txt":
            FileConverter.md_to_txt(input_path, output_path)
        elif conversion_key == "md_to_docx":
            FileConverter.md_to_docx(input_path, output_path)
        elif conversion_key == "md_to_pdf":
            FileConverter.md_to_pdf(input_path, output_path)
        elif conversion_key == "docx_to_txt":
            FileConverter.docx_to_txt(input_path, output_path)
        elif conversion_key == "docx_to_md":
            FileConverter.docx_to_md(input_path, output_path)
        elif conversion_key == "docx_to_pdf":
            FileConverter.docx_to_pdf(input_path, output_path)
        elif conversion_key == "pdf_to_txt":
            FileConverter.pdf_to_txt(input_path, output_path)
        elif conversion_key == "pdf_to_md":
            FileConverter.pdf_to_md(input_path, output_path)
        elif conversion_key == "pdf_to_docx":
            FileConverter.pdf_to_docx(input_path, output_path)
        elif output_format in ['png', 'jpg', 'jpeg'] and input_format == 'pdf':
            # PDF to images - returns list of paths
            os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
            FileConverter.pdf_to_images(input_path, os.path.dirname(output_path) or '.')
        elif input_format in ['png', 'jpg', 'jpeg'] and output_format == 'pdf':
            # Single image to PDF
            FileConverter.images_to_pdf([input_path], output_path)
        else:
            raise ValueError(f"Conversion from {input_format} to {output_format} is not supported")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 conversion_utils.py <input_file> <output_file> [input_format] [output_format]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    input_fmt = sys.argv[3] if len(sys.argv) > 3 else None
    output_fmt = sys.argv[4] if len(sys.argv) > 4 else None
    
    try:
        FileConverter.convert(input_file, output_file, input_fmt, output_fmt)
        print(f"Conversion successful: {input_file} -> {output_file}")
    except Exception as e:
        print(f"Conversion failed: {str(e)}", file=sys.stderr)
        sys.exit(1)
